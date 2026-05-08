"""Audit report & opinion: aggregate engagement signals and build CA-style report sections."""
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional

from openpyxl import Workbook

from app.services.ca_audit_domain import suggest_opinion


async def gather_opinion_signals(db, engagement_id: str) -> Dict[str, Any]:
    """Pull cross-module inputs for opinion and report narrative."""
    mat = await db.ca_materiality.find_one({"engagement_id": engagement_id}, {"_id": 0})
    risks = [r async for r in db.ca_risks.find({"engagement_id": engagement_id}, {"_id": 0}).limit(200)]
    defs = [d async for d in db.ca_control_deficiencies.find({"engagement_id": engagement_id}, {"_id": 0}).limit(200)]
    cases = [c async for c in db.cases.find({"engagement_id": engagement_id, "status": {"$ne": "closed"}}, {"_id": 0}).limit(200)]
    comp = await db.ca_compliance_results.find_one({"engagement_id": engagement_id}, {"_id": 0})
    reqs = (comp or {}).get("requirements") or []
    non_comp = sum(1 for r in reqs if r.get("status") == "non-compliant")
    pending_ev = sum(1 for r in reqs if r.get("status") == "pending evidence")
    findings = [f async for f in db.ca_compliance_findings.find({"engagement_id": engagement_id}, {"_id": 0}).limit(100)]
    snap = await db.ca_fs_snapshots.find_one({"engagement_id": engagement_id}, {"_id": 0}, sort=[("generated_at", -1)])
    fs_issues: List[str] = []
    if snap:
        top_issues = snap.get("issues") if isinstance(snap.get("issues"), list) else []
        fs_issues.extend(str(x) for x in top_issues[:25])
        v = snap.get("validation")
        if isinstance(v, dict):
            for k in ("issues", "warnings", "errors"):
                arr = v.get(k) if isinstance(v.get(k), list) else []
                fs_issues.extend(str(x) for x in arr[:15])
    high_risks = [r for r in risks if r.get("risk_rating") in ("high", "critical")]
    crit_defs = [d for d in defs if (d.get("severity") or "").lower() in ("critical", "high")]
    return {
        "materiality": mat,
        "risk_count": len(risks),
        "high_critical_risks": len(high_risks),
        "deficiency_count": len(defs),
        "critical_high_deficiencies": len(crit_defs),
        "open_cases": len(cases),
        "compliance_non_compliant": non_comp,
        "compliance_pending_evidence": pending_ev,
        "compliance_findings_count": len(findings),
        "fs_snapshot_id": snap.get("id") if snap else None,
        "fs_issue_strings": fs_issues[:25],
        "sample_high_risk_titles": [r.get("title") or r.get("risk_statement") or "Risk" for r in high_risks[:5]],
        "sample_deficiency_titles": [d.get("title") or d.get("description") or "Deficiency" for d in crit_defs[:5]],
    }


def _virtual_observations(signals: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Synthesize observation-shaped rows from modules for opinion logic."""
    v: List[Dict[str, Any]] = []
    if signals.get("compliance_non_compliant", 0) >= 1:
        v.append(
            {
                "title": "Statutory compliance non-compliance",
                "description": f"{signals['compliance_non_compliant']} requirement(s) marked non-compliant.",
                "severity": "high",
                "material": signals["compliance_non_compliant"] >= 2,
                "pervasive": signals["compliance_non_compliant"] >= 5,
                "source": "compliance",
                "resolved": False,
            }
        )
    if signals.get("critical_high_deficiencies", 0) >= 1:
        v.append(
            {
                "title": "Internal control deficiencies",
                "description": f"{signals['critical_high_deficiencies']} high/critical deficiencies open.",
                "severity": "high",
                "material": True,
                "pervasive": signals["critical_high_deficiencies"] >= 3,
                "source": "control",
                "resolved": False,
            }
        )
    if signals.get("open_cases", 0) >= 5:
        v.append(
            {
                "title": "Unresolved audit cases",
                "description": f"{signals['open_cases']} cases remain open.",
                "severity": "medium",
                "material": signals["open_cases"] >= 10,
                "pervasive": False,
                "source": "case",
                "resolved": False,
            }
        )
    fs_n = len(signals.get("fs_issue_strings") or [])
    if fs_n >= 3:
        v.append(
            {
                "title": "Financial statement close / validation issues",
                "description": f"{fs_n} validation signals on latest FS snapshot.",
                "severity": "high",
                "material": fs_n >= 5,
                "pervasive": fs_n >= 12,
                "source": "fs",
                "resolved": False,
            }
        )
    if signals.get("compliance_pending_evidence", 0) >= 8:
        v.append(
            {
                "title": "Evidence scope — compliance checklist",
                "description": "Multiple compliance lines pending evidence — assess scope limitation.",
                "severity": "medium",
                "material": False,
                "pervasive": False,
                "source": "manual",
                "resolved": False,
            }
        )
    return v


def recommend_opinion(observations: List[Dict[str, Any]], signals: Dict[str, Any]) -> Dict[str, Any]:
    """Opinion recommendation: merge module signals with formal observations."""
    merged = list(observations) + _virtual_observations(signals)
    base = suggest_opinion(merged)
    kind = base.get("suggested_opinion") or "unqualified"
    # Disclaimer: insufficient evidence (explicit virtual + pending evidence pressure)
    if signals.get("compliance_pending_evidence", 0) >= 12 and kind == "unqualified":
        kind = "disclaimer"
        base["rationale"] = "Widespread pending-evidence items may indicate scope limitation on statutory assertions."
    if kind == "unqualified":
        display = "clean / unqualified"
    elif kind == "qualified":
        display = "qualified"
    elif kind == "adverse":
        display = "adverse"
    else:
        display = "disclaimer of opinion"
    return {
        **base,
        "suggested_opinion": kind,
        "opinion_display": display,
        "signals_summary": {
            "risks": signals.get("risk_count"),
            "open_cases": signals.get("open_cases"),
            "compliance_nc": signals.get("compliance_non_compliant"),
            "deficiencies_hi_cr": signals.get("critical_high_deficiencies"),
            "fs_issue_count": len(signals.get("fs_issue_strings") or []),
        },
        "virtual_observation_count": len(_virtual_observations(signals)),
    }


def build_report_sections(
    engagement: Dict[str, Any],
    opinion: Optional[Dict[str, Any]],
    observations: List[Dict[str, Any]],
    caro_responses: Optional[Dict[str, Any]],
    signals: Dict[str, Any],
) -> Dict[str, Any]:
    """CA-style report section bodies (draft)."""
    fy = engagement.get("financial_year") or ""
    entity = engagement.get("entity_name") or "the entity"
    op = (opinion or {}).get("suggested_opinion") or "unqualified"
    op_disp = (opinion or {}).get("opinion_display") or op
    kam = [o.get("title") for o in observations if not o.get("resolved")][:8]
    if not kam:
        kam = signals.get("sample_high_risk_titles") or ["No key matters flagged from open observations."]
    mat = signals.get("materiality") or {}
    mat_note = (
        f"Overall materiality was planned at {mat.get('final_materiality')} (performance range "
        f"{mat.get('performance_materiality_low')}–{mat.get('performance_materiality_high')})."
        if mat.get("final_materiality")
        else "Materiality should be documented in the working papers."
    )
    basis = (
        f"We conducted our audit in accordance with Standards on Auditing (SAs) issued by the ICAI, applicable for FY {fy}. "
        f"{mat_note} Those standards require that we plan and perform the audit to obtain reasonable assurance about whether "
        f"the financial statements are free from material misstatement."
    )
    mgmt = (
        "Management is responsible for the preparation and fair presentation of the financial statements in accordance with "
        "the applicable financial reporting framework, and for such internal financial controls as management determines is "
        "necessary to enable preparation of financial statements that are free from material misstatement."
    )
    aud = (
        "Our responsibility is to express an opinion on these financial statements based on our audit. We have taken into "
        "account the provisions of the Act, the accounting and auditing standards and matters which are required to be "
        "included in the audit report under the Act and the Rules made thereunder."
    )
    ifc = (
        "We also address the engagement partner's report on internal financial controls over financial reporting in terms of "
        "the Guidance Note on Audit of IFC under Section 143(3)(i) of the Companies Act, 2013, where applicable."
    )
    caro_annex = caro_responses or {}
    recs: List[str] = []
    if signals.get("compliance_non_compliant", 0):
        recs.append("Remediate open statutory non-compliances and update disclosure checklists.")
    if signals.get("critical_high_deficiencies", 0):
        recs.append("Address control deficiencies and retest remediation.")
    if signals.get("open_cases", 0):
        recs.append("Close or disposition open audit cases with documented responses.")
    if not recs:
        recs.append("Continue monitoring post-reporting matters per audit committee charter.")
    return {
        "opinion": f"Draft opinion classification: {op_disp} ({op}).",
        "basis_for_opinion": basis,
        "key_audit_matters": kam,
        "management_responsibility": mgmt,
        "auditor_responsibility": aud,
        "internal_financial_controls": ifc,
        "caro_annexure": caro_annex.get("responses") if isinstance(caro_annex, dict) else caro_annex,
        "observations": observations,
        "recommendations": recs,
        "entity": entity,
        "financial_year": fy,
    }


def observations_to_xlsx_bytes(observations: List[Dict[str, Any]], engagement_id: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Observations"
    ws.append(["id", "title", "severity", "material", "pervasive", "resolved", "source", "description"])
    for o in observations:
        ws.append(
            [
                o.get("id", ""),
                o.get("title", ""),
                o.get("severity", ""),
                o.get("material", False),
                o.get("pervasive", False),
                o.get("resolved", False),
                o.get("source", ""),
                (o.get("description") or "")[:500],
            ]
        )
    ws.append([])
    ws.append(["engagement_id", engagement_id])
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def report_sections_to_pdf_bytes(sections: Dict[str, Any], engagement_id: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    w, h = A4
    y = h - 50
    c.setFont("Helvetica-Bold", 12)
    c.drawString(50, y, f"Audit Report Draft — {engagement_id}")
    y -= 28
    c.setFont("Helvetica", 9)
    for key in (
        "opinion",
        "basis_for_opinion",
        "management_responsibility",
        "auditor_responsibility",
        "internal_financial_controls",
    ):
        text = sections.get(key) or ""
        c.setFont("Helvetica-Bold", 10)
        c.drawString(50, y, key.replace("_", " ").title())
        y -= 14
        c.setFont("Helvetica", 9)
        for line in str(text).split("\n"):
            while line:
                chunk, line = line[:100], line[100:]
                c.drawString(50, y, chunk)
                y -= 12
                if y < 60:
                    c.showPage()
                    y = h - 50
                    c.setFont("Helvetica", 9)
        y -= 10
    c.setFont("Helvetica-Bold", 10)
    c.drawString(50, y, "Key audit matters")
    y -= 14
    c.setFont("Helvetica", 9)
    for item in sections.get("key_audit_matters") or []:
        c.drawString(50, y, f"· {item}")
        y -= 12
        if y < 60:
            c.showPage()
            y = h - 50
    c.showPage()
    c.save()
    return buf.getvalue()


def report_sections_to_docx_bytes(sections: Dict[str, Any], engagement_id: str) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(f"Audit report draft — {engagement_id}", 0)
    for title, key in [
        ("Opinion", "opinion"),
        ("Basis for opinion", "basis_for_opinion"),
        ("Key audit matters", "key_audit_matters"),
        ("Management responsibility", "management_responsibility"),
        ("Auditor responsibility", "auditor_responsibility"),
        ("Internal financial controls", "internal_financial_controls"),
        ("Recommendations", "recommendations"),
    ]:
        doc.add_heading(title, level=1)
        val = sections.get(key)
        if isinstance(val, list):
            for row in val:
                p = doc.add_paragraph(str(row), style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
        else:
            doc.add_paragraph(str(val or ""))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
